"""
Advanced Resume Editor
Allows you to edit and optimize resumes for multiple job applications.
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Optional
from utils.file_parser import parse_resume, parse_pdf, parse_docx
from utils.groq_optimizer import GroqResumeOptimizer
from docx import Document


class ResumeEditor:
    """Advanced resume editor for job-specific optimizations."""
    
    def __init__(self, groq_api_key: Optional[str] = None):
        self.groq_optimizer = GroqResumeOptimizer(groq_api_key)
        self.model = "llama-3.1-70b-versatile"
    
    def create_job_specific_resume(
        self,
        original_resume_path: str,
        job_description_path: str,
        output_path: str,
        job_title: Optional[str] = None,
        company_name: Optional[str] = None
    ) -> Dict:
        """
        Create a job-specific optimized resume and save it.
        
        Args:
            original_resume_path: Path to original resume
            job_description_path: Path to job description
            output_path: Where to save optimized resume
            job_title: Optional job title for filename
            company_name: Optional company name for filename
        
        Returns:
            Dictionary with results and metadata
        """
        print(f"📄 Processing resume for: {job_title or 'Job Application'}")
        
        # Load files
        resume_text = parse_resume(original_resume_path)
        if not resume_text:
            return {"error": "Could not parse resume"}
        
        job_description = parse_resume(job_description_path)
        if not job_description:
            return {"error": "Could not parse job description"}
        
        if not self.groq_optimizer.is_available():
            return {"error": "Groq API not available"}
        
        # Generate optimized resume
        print("🤖 Generating optimized resume with Groq...")
        result = self.groq_optimizer.create_optimized_resume(
            resume_text, job_description, self.model
        )
        
        if "error" in result:
            return result
        
        optimized_text = result.get("optimized_resume", "")
        
        # Save optimized resume
        try:
            # Determine output format based on extension
            if output_path.endswith('.docx'):
                self._save_as_docx(optimized_text, output_path)
            elif output_path.endswith('.pdf'):
                # For PDF, save as text first (PDF generation requires additional libraries)
                txt_path = output_path.replace('.pdf', '.txt')
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(optimized_text)
                print(f"⚠️  PDF format not directly supported. Saved as: {txt_path}")
                output_path = txt_path
            else:
                # Save as text
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(optimized_text)
            
            # Create metadata file
            metadata = {
                "original_resume": original_resume_path,
                "job_description": job_description_path,
                "output_path": output_path,
                "job_title": job_title,
                "company_name": company_name,
                "created_at": datetime.now().isoformat(),
                "model_used": self.model
            }
            
            metadata_path = output_path.replace('.txt', '_metadata.json').replace('.docx', '_metadata.json')
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2)
            
            return {
                "success": True,
                "output_path": output_path,
                "metadata_path": metadata_path,
                "metadata": metadata
            }
        
        except Exception as e:
            return {"error": f"Error saving resume: {str(e)}"}
    
    def batch_optimize(
        self,
        original_resume_path: str,
        job_descriptions_dir: str,
        output_dir: str,
        job_info_file: Optional[str] = None
    ) -> Dict:
        """
        Optimize resume for multiple job applications.
        
        Args:
            original_resume_path: Path to original resume
            job_descriptions_dir: Directory containing job description files
            output_dir: Directory to save optimized resumes
            job_info_file: Optional JSON file with job titles and company names
        
        Returns:
            Dictionary with results for each job
        """
        import glob
        
        # Load job info if provided
        job_info = {}
        if job_info_file and os.path.exists(job_info_file):
            with open(job_info_file, 'r', encoding='utf-8') as f:
                job_info = json.load(f)
        
        # Find all job description files
        job_files = glob.glob(os.path.join(job_descriptions_dir, "*.txt"))
        job_files.extend(glob.glob(os.path.join(job_descriptions_dir, "*.pdf")))
        
        if not job_files:
            return {"error": f"No job description files found in {job_descriptions_dir}"}
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        results = []
        
        print(f"\n🚀 Processing {len(job_files)} job applications...\n")
        
        for i, job_file in enumerate(job_files, 1):
            job_name = os.path.splitext(os.path.basename(job_file))[0]
            job_title = job_info.get(job_name, {}).get("title", job_name)
            company_name = job_info.get(job_name, {}).get("company", "Unknown")
            
            print(f"[{i}/{len(job_files)}] Processing: {job_title} at {company_name}")
            
            # Create output filename
            safe_job_title = "".join(c for c in job_title if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
            safe_company = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()[:30]
            output_filename = f"resume_{safe_company}_{safe_job_title}.txt"
            output_path = os.path.join(output_dir, output_filename)
            
            # Optimize
            result = self.create_job_specific_resume(
                original_resume_path,
                job_file,
                output_path,
                job_title,
                company_name
            )
            
            result["job_file"] = job_file
            result["job_title"] = job_title
            result["company_name"] = company_name
            results.append(result)
            
            if "error" in result:
                print(f"  ❌ Error: {result['error']}")
            else:
                print(f"  ✅ Saved: {output_path}")
        
        # Create summary
        summary = {
            "total_jobs": len(job_files),
            "successful": len([r for r in results if "success" in r]),
            "failed": len([r for r in results if "error" in r]),
            "results": results,
            "created_at": datetime.now().isoformat()
        }
        
        summary_path = os.path.join(output_dir, "batch_summary.json")
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)
        
        print(f"\n📊 Summary saved to: {summary_path}")
        print(f"   Successful: {summary['successful']}/{summary['total_jobs']}")
        
        return summary
    
    def _save_as_docx(self, text: str, output_path: str):
        """Save text as DOCX file."""
        doc = Document()
        
        # Split into paragraphs
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # Check if it's a heading (all caps or short line)
                if para.isupper() or (len(para) < 50 and '\n' not in para):
                    doc.add_heading(para.strip(), level=2)
                else:
                    doc.add_paragraph(para.strip())
        
        doc.save(output_path)
    
    def compare_versions(
        self,
        original_path: str,
        optimized_path: str
    ) -> Dict:
        """Compare original and optimized resume versions."""
        original = parse_resume(original_path)
        optimized = parse_resume(optimized_path)
        
        if not original or not optimized:
            return {"error": "Could not parse one or both files"}
        
        # Basic comparison
        original_words = set(original.lower().split())
        optimized_words = set(optimized.lower().split())
        
        added_words = optimized_words - original_words
        removed_words = original_words - optimized_words
        
        return {
            "original_word_count": len(original.split()),
            "optimized_word_count": len(optimized.split()),
            "words_added": len(added_words),
            "words_removed": len(removed_words),
            "sample_added_words": list(added_words)[:20],
            "sample_removed_words": list(removed_words)[:20]
        }


def main():
    """Command-line interface for resume editor."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Resume Editor - Create job-specific optimized resumes',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Optimize for single job
  python resume_editor.py -r resume.pdf -j job.txt -o optimized_resume.txt --job-title "Software Engineer" --company "Tech Corp"
  
  # Batch optimize for multiple jobs
  python resume_editor.py --batch -r resume.pdf --jobs-dir ./job_descriptions --output-dir ./optimized_resumes
        """
    )
    
    parser.add_argument('-r', '--resume', required=True, help='Original resume file')
    parser.add_argument('-j', '--job', help='Job description file (single job mode)')
    parser.add_argument('-o', '--output', help='Output file path (single job mode)')
    parser.add_argument('--job-title', help='Job title')
    parser.add_argument('--company', help='Company name')
    parser.add_argument('--batch', action='store_true', help='Batch mode')
    parser.add_argument('--jobs-dir', help='Directory with job descriptions (batch mode)')
    parser.add_argument('--output-dir', help='Output directory (batch mode)')
    parser.add_argument('--job-info', help='JSON file with job titles and companies')
    parser.add_argument('--api-key', help='Groq API key')
    
    args = parser.parse_args()
    
    # Set API key
    if args.api_key:
        os.environ['GROQ_API_KEY'] = args.api_key
    
    if not os.getenv('GROQ_API_KEY'):
        print("❌ Error: GROQ_API_KEY not set!")
        print("   Set it as environment variable or use --api-key")
        sys.exit(1)
    
    editor = ResumeEditor()
    
    if args.batch:
        # Batch mode
        if not args.jobs_dir or not args.output_dir:
            print("❌ Error: --jobs-dir and --output-dir required for batch mode")
            sys.exit(1)
        
        results = editor.batch_optimize(
            args.resume,
            args.jobs_dir,
            args.output_dir,
            args.job_info
        )
        
        if "error" in results:
            print(f"❌ Error: {results['error']}")
            sys.exit(1)
    else:
        # Single job mode
        if not args.job or not args.output:
            print("❌ Error: --job and --output required for single job mode")
            sys.exit(1)
        
        result = editor.create_job_specific_resume(
            args.resume,
            args.job,
            args.output,
            args.job_title,
            args.company
        )
        
        if "error" in result:
            print(f"❌ Error: {result['error']}")
            sys.exit(1)
        
        print(f"\n✅ Optimized resume saved to: {result['output_path']}")


if __name__ == "__main__":
    import sys
    main()

