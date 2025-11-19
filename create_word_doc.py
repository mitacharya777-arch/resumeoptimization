"""
Script to create Word document from complete documentation.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_word_documentation():
    """Create Word document from complete documentation."""
    
    # Read the markdown file
    with open('COMPLETE_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Resume Optimizer', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Complete Project Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata paragraph
    meta = doc.add_paragraph()
    meta.add_run('Version: 1.0\n').bold = True
    meta.add_run('Last Updated: 2024\n')
    meta.add_run('Comprehensive Guide to Resume Optimizer Project')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Parse and add content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines at start
        if not line and i == 0:
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('```'):
            # Code block - skip the ``` line
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                p.add_run('\n'.join(code_lines)).font.name = 'Courier New'
                p.add_run('\n').font.name = 'Courier New'
        elif line.startswith('`') and line.endswith('`'):
            # Inline code
            p = doc.add_paragraph()
            p.add_run(line[1:-1]).font.name = 'Courier New'
        elif line.startswith('|'):
            # Table - skip for now (complex to parse)
            pass
        elif line.startswith('=') and all(c == '=' for c in line):
            # Separator - add paragraph
            doc.add_paragraph('─' * 50)
        elif line:
            # Regular paragraph
            doc.add_paragraph(line)
        else:
            # Empty line - add spacing
            doc.add_paragraph()
        
        i += 1
    
    # Add footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.add_run('End of Documentation').bold = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    output_file = 'Resume_Optimizer_Complete_Documentation.docx'
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    return output_file

if __name__ == '__main__':
    try:
        create_word_documentation()
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")
        print("Make sure python-docx is installed: pip install python-docx")

