"""
Convert TECHNICAL_DOCUMENTATION.md to Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_technical_docx():
    """Create Word document from technical documentation."""
    
    # Read markdown file
    with open('TECHNICAL_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Technical Documentation: Resume Optimization Application', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata paragraph
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run('Comprehensive Technical Documentation')
    meta_run.font.size = Pt(12)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Parse content
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.style.font.color.rgb = RGBColor(104, 42, 83)  # Maroon
            i += 1
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.style.font.color.rgb = RGBColor(104, 42, 83)  # Maroon
            i += 1
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.style.font.color.rgb = RGBColor(104, 42, 83)  # Maroon
            i += 1
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.style.font.color.rgb = RGBColor(104, 42, 83)  # Maroon
            i += 1
        # Code blocks
        elif line.startswith('```'):
            # Skip code block marker
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip closing ```
            
            # Add code block as formatted text
            if code_lines:
                para = doc.add_paragraph()
                para.style = 'No Spacing'
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green for code
                para_format = para.paragraph_format
                para_format.left_indent = Inches(0.5)
                para_format.space_before = Pt(6)
                para_format.space_after = Pt(6)
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            para = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
            i += 1
        # Bold text (markdown **text**)
        elif '**' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
            i += 1
        # Inline code (`code`)
        elif '`' in line:
            para = doc.add_paragraph()
            parts = re.split(r'(`[^`]+`)', line)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                else:
                    para.add_run(part)
            i += 1
        # Horizontal rule
        elif line.startswith('---'):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(12)
            para_format.space_after = Pt(12)
            # Add a border
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
        # Regular text
        else:
            # Check for special formatting
            if line.startswith('**') and line.endswith('**'):
                para = doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
                run.font.size = Pt(12)
            elif line.startswith('✅') or line.startswith('⚠️') or line.startswith('🔍'):
                para = doc.add_paragraph(line)
                para_format = para.paragraph_format
                para_format.space_before = Pt(6)
            else:
                para = doc.add_paragraph(line)
            i += 1
    
    # Add footer with document info
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Technical Documentation: Resume Optimization Application | Generated from TECHNICAL_DOCUMENTATION.md"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    output_file = 'TECHNICAL_DOCUMENTATION.docx'
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    return output_file

if __name__ == '__main__':
    try:
        create_technical_docx()
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

