"""
Document Generator for Resume Downloads
Generates formatted PDF and Word documents from resume text.
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from xml.sax.saxutils import escape
import re


def generate_pdf(resume_text: str) -> BytesIO:
    """
    Generate a formatted PDF from resume text.
    Professional, ATS-friendly format with proper formatting.

    Args:
        resume_text: Plain text resume content

    Returns:
        BytesIO: PDF file in memory
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )

    story = []
    styles = getSampleStyleSheet()

    # Professional styles with good spacing
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        leading=20
    )

    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=16,
        alignment=TA_CENTER,
        fontName='Helvetica',
        leading=14
    )

    section_style = ParagraphStyle(
        'Section',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=14,
        fontName='Helvetica-Bold',
        leading=16
    )

    # Bold style for company/job headers
    company_style = ParagraphStyle(
        'Company',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=4,
        spaceBefore=6,
        fontName='Helvetica-Bold',
        leading=14
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=5,
        fontName='Helvetica',
        leading=13
    )

    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        leftIndent=20,
        fontName='Helvetica',
        leading=13
    )

    # Section keywords
    section_keywords = [
        'SUMMARY', 'PROFESSIONAL SUMMARY', 'EXPERIENCE', 'WORK EXPERIENCE',
        'EDUCATION', 'SKILLS', 'TECHNICAL SKILLS', 'PROJECTS',
        'CERTIFICATIONS', 'ACHIEVEMENTS', 'AWARDS'
    ]

    # Parse resume line by line
    lines = resume_text.strip().split('\n')
    line_count = 0
    in_experience_section = False

    for i, line in enumerate(lines):
        line = line.strip()

        if not line:
            story.append(Spacer(1, 0.12*inch))
            continue

        # Escape special XML characters
        line_escaped = escape(line)

        # First line is name
        if line_count == 0:
            story.append(Paragraph(line_escaped, name_style))
            line_count += 1
            continue

        # Second/third lines might be title and contact
        if line_count < 3 and not any(keyword in line.upper() for keyword in section_keywords):
            story.append(Paragraph(line_escaped, contact_style))
            line_count += 1
            continue

        # Check if it's a section header
        is_section = line.upper() in section_keywords

        if is_section:
            in_experience_section = ('EXPERIENCE' in line.upper())
            story.append(Spacer(1, 0.18*inch))
            story.append(Paragraph(f"<b>{line_escaped}</b>", section_style))

        # Check if it's a company/job header (contains pipe separator)
        elif '|' in line and in_experience_section:
            # This is a company header - make it bold
            story.append(Paragraph(f"<b>{line_escaped}</b>", company_style))

        # Check for project names (ALL CAPS lines that are not section headers)
        elif line.isupper() and len(line) > 3 and not is_section and not '|' in line:
            # This is likely a project name - make it bold
            story.append(Paragraph(f"<b>{line_escaped}</b>", company_style))

        elif line.startswith('-') or line.startswith('•'):
            # Bullet point
            bullet_text = line[1:].strip()
            bullet_escaped = escape(bullet_text)
            story.append(Paragraph(f"• {bullet_escaped}", bullet_style))
        else:
            # Regular line
            story.append(Paragraph(line_escaped, body_style))

    # Build PDF
    try:
        doc.build(story)
    except Exception as e:
        # If build fails, create a simple error page
        story = [Paragraph(f"Error generating PDF: {str(e)}", body_style)]
        doc.build(story)

    buffer.seek(0)
    return buffer


def generate_docx(resume_text: str) -> BytesIO:
    """
    Generate a formatted Word document from resume text.
    Professional, ATS-friendly format with proper formatting.

    Args:
        resume_text: Plain text resume content

    Returns:
        BytesIO: DOCX file in memory
    """
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Section keywords
    section_keywords = [
        'SUMMARY', 'PROFESSIONAL SUMMARY', 'EXPERIENCE', 'WORK EXPERIENCE',
        'EDUCATION', 'SKILLS', 'TECHNICAL SKILLS', 'PROJECTS',
        'CERTIFICATIONS', 'ACHIEVEMENTS', 'AWARDS'
    ]

    # Parse resume line by line
    lines = resume_text.strip().split('\n')
    line_count = 0
    in_experience_section = False

    for line in lines:
        line = line.strip()

        if not line:
            # Add small spacing
            p = doc.add_paragraph()
            p.space_after = Pt(6)
            continue

        # First line is name
        if line_count == 0:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            p.space_after = Pt(4)
            line_count += 1
            continue

        # Second/third lines might be title and contact
        if line_count < 3 and not any(keyword in line.upper() for keyword in section_keywords):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(10)
            if line_count == 2:
                p.space_after = Pt(12)
            line_count += 1
            continue

        # Check if it's a section header
        is_section = line.upper() in section_keywords

        if is_section:
            in_experience_section = ('EXPERIENCE' in line.upper())
            # Add spacing before section
            spacer = doc.add_paragraph()
            spacer.space_after = Pt(4)

            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            p.space_after = Pt(6)

        # Check if it's a company/job header (contains pipe separator)
        elif '|' in line and in_experience_section:
            # This is a company header - make it bold
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(11)
            run.font.bold = True
            p.space_before = Pt(6)
            p.space_after = Pt(3)

        # Check for project names (ALL CAPS lines that are not section headers)
        elif line.isupper() and len(line) > 3 and not is_section and not '|' in line:
            # This is likely a project name - make it bold
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(11)
            run.font.bold = True
            p.space_before = Pt(6)
            p.space_after = Pt(3)

        elif line.startswith('-') or line.startswith('•'):
            # Bullet point
            bullet_text = line[1:].strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            run = p.runs[0]
            run.font.size = Pt(10)
            p.space_after = Pt(2)
        else:
            # Regular text
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(10)
            p.space_after = Pt(4)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
